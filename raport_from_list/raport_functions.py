from users.models import Profile, ModifiedFile, PDFFile, ExcelFile, DocxZipFile, PdfZipFile
from docx import Document
import tempfile
import zipfile
from io import BytesIO
from django.core.files.base import ContentFile
import os
from bs4 import BeautifulSoup
import re
import subprocess

def convert_docxs_to_zip(profile, modified_files):
    """
    Converts a list of ModifiedFile objects into a zip file.

    Args:
    profile (Profile): The user's profile object.
    modified_files (QuerySet): A queryset of ModifiedFile objects.

    Returns:
    str: URL to the created zip file.
    """
    try:
        # Check if there is an existing zip file
        existing_zip_file = DocxZipFile.objects.filter(profile=profile).first()
        if existing_zip_file:
            return existing_zip_file.docx_zip_file.url

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            for modified_file in modified_files:
                file_path = os.path.join(modified_file.modyfied_doc_file.path)
                zip_file.write(file_path, os.path.basename(file_path))

        zip_buffer.seek(0)
        new_zip_file = DocxZipFile(profile=profile)
        new_zip_file_name = 'all_docx_files.zip'
        new_zip_file.docx_zip_file.save(new_zip_file_name, ContentFile(zip_buffer.getvalue()))
        new_zip_file.save()

        return new_zip_file.docx_zip_file.url
    except Exception as e:
        # Handle the exception or log the error
        print(f"Error generating zip from docx: {str(e)}")
        return None


def convert_to_pdf(modified_files, username):
    """
    Converts a list of ModifiedFile objects to PDF format.

    Args:
    modified_files (QuerySet): A queryset of ModifiedFile objects.
    username (str): The username of the profile.

    Returns:
    list: URLs of the converted PDF files.
    """
    profile = Profile.objects.get(user__username=username)
    old_pdf_files = PDFFile.objects.filter(profile=profile)

    if old_pdf_files:
        # PDF files already exist, no need to perform conversion
        pdf_urls = [pdf_file.pdf_file.url for pdf_file in old_pdf_files]
    else:
        try:
            pdf_urls = []
            doc_dir_path = os.path.dirname(modified_files[0].modyfied_doc_file.path)

            for file_name in os.listdir(doc_dir_path):
                if file_name.endswith(".docx"):
                    docx_file_path = os.path.join(doc_dir_path, file_name)
                    pdf_file_name = os.path.splitext(file_name)[0] + ".pdf"
                    pdf_file_path = os.path.join(doc_dir_path, pdf_file_name)

                    # Execute doc2pdf command-line tool
                    subprocess.run(["/usr/bin/doc2pdf", docx_file_path, pdf_file_path])

                    # Save PDF file information to the database
                    relative_file_path = os.path.join(username, 'modified_doc_documents', pdf_file_name)
                    pdf_file = PDFFile.objects.create(pdf_file=relative_file_path, profile=profile)
                    pdf_file.save()
                    pdf_urls.append(pdf_file.pdf_file.url)

            # Print the PDF file names for testing purposes
            pdfs = PDFFile.objects.filter(profile=profile)
            for pdf in pdfs:
                print(f"PDF file: {pdf.pdf_file.name}")

        except Exception as e:
            print(f"Error converting to PDF: {str(e)}")
            pdf_urls = []

    return pdf_urls


def convert_pdfs_to_zip(profile):
    """
    Converts all PDF files associated with a profile into a single zip file.

    Args:
    profile (Profile): The user's profile object.

    Returns:
    str: URL to the created zip file.
    """
    try:
        # Check if there is an existing zip file
        existing_zip_file = PdfZipFile.objects.filter(profile=profile).first()
        if existing_zip_file:
            return existing_zip_file.pdf_zip_file.url

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
            pdf_files = PDFFile.objects.filter(profile=profile)
            for pdf_file in pdf_files:
                file_path = os.path.join(pdf_file.pdf_file.path)
                zip_file.write(file_path, os.path.basename(file_path))

        zip_buffer.seek(0)
        new_zip_file = PdfZipFile(profile=profile)
        new_zip_file_name = 'all_pdf_files.zip'
        new_zip_file.pdf_zip_file.save(new_zip_file_name, ContentFile(zip_buffer.getvalue()))
        new_zip_file.save()

        return new_zip_file.pdf_zip_file.url
    except Exception as e:
        # Handle the exception or log the error
        print(f"Error generating PDF zip file: {str(e)}")
        return None

def delete_old_files(profile):
    """
    Deletes all files associated with a profile, including ModifiedFile, 
    PDFFile, ExcelFile, DocxZipFile, and PdfZipFile objects.

    Args:
    profile (Profile): The user's profile object.
    """
    try:
        # Delete ModifiedFile objects
        ModifiedFile.objects.filter(profile=profile).delete()

        # Delete PDFFile objects
        PDFFile.objects.filter(profile=profile).delete()

        # Delete ExcelFile objects
        ExcelFile.objects.filter(profile=profile).delete()

        # Delete DocxZipFile objects
        DocxZipFile.objects.filter(profile=profile).delete()

        # Delete PdfZipFile objects
        PdfZipFile.objects.filter(profile=profile).delete()

        # Delete doc_file from Profile
        if profile.doc_file:
            profile.doc_file.delete()
            profile.doc_file = None
            profile.save()

        # Delete other model objects if needed

    except Exception as e:
        # Handle the exception or log the error
        print(f"Error deleting old files: {str(e)}")

def generate_modified_filename(prefix, file_number):
    modified_filename = f"{prefix}_{file_number}.docx"
    return modified_filename



def generate_reports(profile, form_data, request):
    """
    Generates modified reports based on form data and original file. 
    Replaces specific words in the document as per Excel data.

    Args:
    profile (Profile): The user's profile object.
    form_data (dict): Data from the form mapping Excel column values to words.
    request (HttpRequest): The request object.

    Returns:
    dict: A dictionary containing URLs of modified documents and any words not found.
    """
    try:
        original_file_path = profile.doc_file.path
        docx_urls = []
        not_found = []
        words_number = len(form_data[list(form_data.keys())[0]])
        selected_filenames = request.session.get('selected_filenames', None)
        print(f"to jest przekazany słownik {form_data}")
        for i in range(words_number):
            with open(original_file_path, "rb") as original_file:
                doc = Document(original_file)
            print("otwarcie documentu original_file_path")

            for search_key in form_data.keys():
                found = False

                for paragraph in doc.paragraphs:
                    if search_key in paragraph.text:
                        found = True
                        paragraph.text = paragraph.text.replace(search_key, form_data[search_key][i])

            if not found:
                not_found.append(search_key)
            print(f"Zamiana {search_key} na {form_data[search_key][i]}")

            if selected_filenames:
                modified_filename = selected_filenames[i]
            else:
                modified_filename = generate_modified_filename("file", i)

            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                doc.save(temp_file.name)
                temp_file.close()

            with open(temp_file.name, "rb") as file:
                modified_file = ModifiedFile(profile=profile)
                modified_file.modyfied_doc_file.save(modified_filename, file, save=False)
                modified_file.modyfied_doc_file_name = modified_filename
                modified_file.save()

                docx_urls.append(modified_file.modyfied_doc_file.url)

            os.unlink(temp_file.name)
            print("zapisanie dokumentu")

        return {
            'docx_urls': docx_urls,
            'not_found': not_found,
        }
    except Exception as e:
        # Handle the exception or log the error
        print(f"Error generating reports: {str(e)}")
        return {
            'docx_urls': [],
            'not_found': [],
        }

  
def generate_html(profile):
    """
    Generates an HTML representation of a DOCX file's content.

    Args:
    profile (Profile): The user's profile object.

    Returns:
    str: HTML representation of the DOCX file.
    """

    original_file_path = profile.doc_file.path

    with open(original_file_path, "rb") as original_file:
        document = Document(original_file)
        soup = BeautifulSoup('', 'html.parser')

        for para in document.paragraphs:
            p = soup.new_tag('p')
            for run in para.runs:
                words = re.split(r'(\s+)', run.text)  
                for word in words:
                    if word.strip():  
                        span = soup.new_tag('span')
                        if run.bold:
                            span['style'] = 'font-weight: bold;'
                        if run.italic:
                            span['style'] = span.get('style', '') + 'font-style: italic;'
                        if run.underline:
                            span['style'] = span.get('style', '') + 'text-decoration: underline;'
                        span.string = word
                        p.append(span)
                    else: 
                        p.append(word)  
            soup.append(p)

        return str(soup)